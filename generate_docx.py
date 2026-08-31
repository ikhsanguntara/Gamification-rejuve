import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Set inner margins (padding) of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_rejuve_doc():
    doc = Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Palette Colors
    COLOR_PRIMARY = RGBColor(131, 24, 67)     # Re.juve Burgundy (#831843)
    COLOR_SECONDARY = RGBColor(157, 23, 77)   # Deep Rose (#9d174d)
    COLOR_TEXT_MAIN = RGBColor(30, 41, 59)    # Slate 800 (#1e293b)
    COLOR_TEXT_MUTED = RGBColor(100, 116, 139)# Slate 500 (#64748b)
    COLOR_BG_HEADER = "831843"                # Hex Burgundy
    COLOR_BG_LIGHT = "FDF2F8"                 # Light Pink / Rose
    COLOR_BG_GRAY = "F8FAFC"                  # Light Slate

    # Set normal style font
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT_MAIN

    # ==================== COVER / HEADER ====================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    r_brand = p_title.add_run("RE.JUVE TRUE COLD-PRESSED\n")
    r_brand.font.size = Pt(14)
    r_brand.font.bold = True
    r_brand.font.color.rgb = COLOR_SECONDARY

    r_title = p_title.add_run("Dokumen Spesifikasi Sistem & Alur Operasional\n")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    r_sub = p_title.add_run("Store Performance, Multi-Tier Approval, & Crew Gamification Management")
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = COLOR_TEXT_MUTED

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(24)
    r_meta = p_meta.add_run("Versi 2.0 (Production Ready) • Panduan Materi Presentasi Manajemen")
    r_meta.font.size = Pt(10)
    r_meta.font.color.rgb = COLOR_TEXT_MUTED

    # Horizontal Divider Line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(16)
    r_div = p_div.add_run("―" * 55)
    r_div.font.color.rgb = RGBColor(226, 232, 240)
    r_div.font.size = Pt(14)

    # ==================== 1. EXECUTIVE SUMMARY ====================
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Ringkasan Eksekutif (Executive Summary)")
    r_h1.font.color.rgb = COLOR_PRIMARY
    r_h1.font.bold = True

    p = doc.add_paragraph(
        "Sistem Re.juve Store Gamification & Mission Management adalah platform digital operasional berbasis web/mobile "
        "yang dirancang untuk mentransformasi kepatuhan Standard Operating Procedure (SOP) gerai Re.juve—mencakup "
        "#CleanLabel, standar Cold Chain (2–4°C), sanitasi higienis, dan kecepatan pelayanan (< 45 detik)—dari metode manual "
        "menjadi ekosistem misi terstruktur berhadiah bintang (gamified store missions)."
    )
    p.paragraph_format.space_after = Pt(8)

    p_val_intro = doc.add_paragraph()
    r_vi = p_val_intro.add_run("Nilai Tambah Utama (Key Value Propositions):")
    r_vi.bold = True

    bullets = [
        ("Transparansi & Akuntabilitas Bertingkat (Two-Tier Quality Gate): ", 
         "Memisahkan wewenang antara Area Supervisor sebagai penilai lapangan dan Head of Operations sebagai pengambil keputusan akhir."),
        ("Isolasi Data Penuh per Batch Gerai: ", 
         "Supervisor dan Head hanya dapat melihat dan mengevaluasi gerai/batch yang ditugaskan kepada mereka, mencegah kebocoran data antar cabang."),
        ("Reward Engine Terintegrasi Otomatis: ", 
         "Pencairan bintang (+5 Stars per kru), akumulasi XP, kenaikan level karir, dan pembukaan badge terjadi secara otomatis begitu evaluasi disetujui Head."),
        ("Master Template SOP Modular: ", 
         "Superadmin dapat mengelola paket template SOP (Standar Mall, Kiosk Express, Onboarding Kru) dan menerapkannya ke batch baru hanya dalam 1 klik.")
    ]
    for b_title, b_desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r_bt = bp.add_run(b_title)
        r_bt.bold = True
        r_bt.font.color.rgb = COLOR_SECONDARY
        bp.add_run(b_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ==================== 2. HIERARKI PERAN & RBAC ====================
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. Hierarki Peran & Matriks Wewenang (RBAC)")
    r_h2.font.color.rgb = COLOR_PRIMARY
    r_h2.font.bold = True

    p = doc.add_paragraph(
        "Sistem memiliki 4 persona utama dengan wewenang yang saling terintegrasi dan terisolasi secara aman:"
    )
    p.paragraph_format.space_after = Pt(8)

    roles_data = [
        ("1. Store Crew (Store Specialist)", "Pelaksana SOP di gerai. Memantau progres misi mingguan, mengumpulkan reward bintang, menaikkan level keahlian (Level 1–10), dan memantau peringkat leaderboard gerai."),
        ("2. Area Store Supervisor (Field Evaluator)", "Auditor operasional lapangan. Bertugas mengunjungi gerai, menilai setiap butir kriteria SOP, memberi skor per kru, melampirkan foto bukti audit (logbook/chiller), menyimpan draft, dan mengirim ke Head."),
        ("3. Head of Operations (Quality Approver)", "Pintu kendali mutu akhir. Meninjau kelayakan bukti foto dan skor evaluasi. Memiliki hak penuh untuk APPROVE (mencairkan bintang serentak) atau REQUEST REVISION (mengembalikan ke SPV dengan catatan wajib)."),
        ("4. System Superadmin (Master Controller)", "Pengendali sistem nasional. Mengelola master batch gerai, mengonfigurasi katalog master template SOP, menugaskan supervisor & head, serta mengelola direktori master pengguna.")
    ]
    for r_title, r_desc in roles_data:
        rp = doc.add_paragraph(style='List Bullet')
        rp.paragraph_format.space_after = Pt(4)
        r_rt = rp.add_run(f"{r_title}: ")
        r_rt.bold = True
        r_rt.font.color.rgb = COLOR_PRIMARY
        rp.add_run(r_desc)

    # Table of RBAC Matrix
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    p_tbl_title = doc.add_paragraph()
    r_tt = p_tbl_title.add_run("Tabel Matriks Hak Akses Modul Sistem:")
    r_tt.bold = True

    table = doc.add_table(rows=10, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Modul / Fitur Sistem", "Store Crew", "Supervisor", "Head of Ops", "Superadmin"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], COLOR_BG_HEADER)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)

    matrix_rows = [
        ("Melihat Misi & SOP Gerai", "Ya (Batch Sendiri)", "Ya (Batch Ditugaskan)", "Ya (Batch Ditugaskan)", "Ya (Semua Batch)"),
        ("Melihat Bintang, Level & Journey", "Ya", "Ya", "Ya", "Ya"),
        ("Input Skor & Upload Bukti Audit", "Tidak", "Ya", "Tidak", "Tidak"),
        ("Simpan Draft & Submit ke Head", "Tidak", "Ya", "Tidak", "Tidak"),
        ("Persetujuan Akhir (Approve/Revise)", "Tidak", "Tidak", "Ya", "Ya"),
        ("Pemberian Reward Bintang Otomatis", "Penerima", "Tidak", "Pemicu (Trigger)", "Ya"),
        ("Kelola Master Template SOP", "Tidak", "Tidak", "Tidak", "Ya"),
        ("Buat / Edit / Hapus Batch Gerai", "Tidak", "Tidak", "Tidak", "Ya"),
        ("Manajemen Pengguna & Roster Tim", "Tidak", "Tidak", "Tidak", "Ya")
    ]

    for row_idx, data in enumerate(matrix_rows, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = COLOR_BG_GRAY if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(9)
                if text == "Ya":
                    r.font.color.rgb = RGBColor(16, 185, 129) # Emerald
                    r.font.bold = True
                elif text == "Tidak":
                    r.font.color.rgb = RGBColor(148, 163, 184) # Gray
                elif "Pemicu" in text or "Penerima" in text:
                    r.font.color.rgb = COLOR_PRIMARY
                    r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== 3. ALUR OPERASIONAL SISTEM ====================
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Diagram & Penjelasan Alur Operasional (System Flows)")
    r_h3.font.color.rgb = COLOR_PRIMARY
    r_h3.font.bold = True

    # 3.1 Siklus 3 Minggu
    h3_1 = doc.add_heading(level=2)
    r_h3_1 = h3_1.add_run("3.1 Siklus 3 Minggu Operasional Gerai (Batch Lifecycle)")
    r_h3_1.font.color.rgb = COLOR_SECONDARY

    p = doc.add_paragraph(
        "Setiap batch gerai menjalankan siklus 3 minggu terstruktur yang mencakup seluruh aspek operasional:"
    )
    p.paragraph_format.space_after = Pt(4)

    weeks_info = [
        ("Week 1 - Fondasi Higienitas & Rantai Dingin: ", "Fokus pada audit suhu chiller display & walk-in (2–4°C), sanitasi mesin press buah, penggunaan sarung tangan/masker, dan kalibrasi termometer digital."),
        ("Week 2 - Kualitas Resep & Kecepatan Layanan: ", "Fokus pada uji kemanisan alami buah (Brix measurement), kepatuhan resep tanpa gula/air tambahan (#CleanLabel), dan speed of service barista (< 45 detik saat antrean)."),
        ("Week 3 - Audit Akhir, Stok & Penutupan Siklus: ", "Fokus pada stock opname buah segar, pemeriksaan mesin kasir POS/EDC, prosedur closing gerai, dan evaluasi kelulusan batch untuk penentuan gerai terbaik.")
    ]
    for w_title, w_desc in weeks_info:
        wp = doc.add_paragraph(style='List Bullet')
        wp.paragraph_format.space_after = Pt(4)
        r_wt = wp.add_run(w_title)
        r_wt.bold = True
        wp.add_run(w_desc)

    # 3.2 Alur Evaluasi Supervisor
    h3_2 = doc.add_heading(level=2)
    r_h3_2 = h3_2.add_run("3.2 Alur Penilaian Lapangan oleh Supervisor")
    r_h3_2.font.color.rgb = COLOR_SECONDARY

    p_spv_steps = [
        "1. Supervisor mengunjungi gerai dan membuka halaman Evaluasi di perangkat tablet/smartphone.",
        "2. Sistem secara otomatis menyaring batch yang ditugaskan kepada Supervisor tersebut.",
        "3. Supervisor memilih misi yang sedang aktif pada minggu berjalan (misal: 'Cek Suhu Chiller 2-4°C').",
        "4. Supervisor memeriksa butir kriteria SOP dan memasukkan skor (0–100) untuk masing-masing kru yang bertugas.",
        "5. Supervisor wajib mengambil & mengunggah minimal 1 foto bukti observasi (contoh: foto termometer/logbook).",
        "6. Supervisor dapat memilih 'Simpan Draft' atau langsung 'Kirim ke Head of Operations' (Status berubah menjadi PENDING_REVIEW)."
    ]
    for step in p_spv_steps:
        p = doc.add_paragraph(step)
        p.paragraph_format.space_after = Pt(3)

    # 3.3 Alur Persetujuan Head
    h3_3 = doc.add_heading(level=2)
    r_h3_3 = h3_3.add_run("3.3 Alur Persetujuan & Pencairan Reward oleh Head of Operations")
    r_h3_3.font.color.rgb = COLOR_SECONDARY

    p_head_steps = [
        "1. Head of Operations membuka halaman Antrean Persetujuan (/approvals).",
        "2. Sistem hanya menampilkan antrean evaluasi dari cabang yang berada di bawah wewenang Head tersebut.",
        "3. Head memeriksa lembar penilaian, nilai rata-rata, komentar supervisor, serta foto bukti pendukung.",
        "4. Opsi A - APPROVE (Disetujui):",
        "   • Status misi berubah menjadi COMPLETED / APPROVED.",
        "   • Sistem secara otomatis mencairkan bintang (+5 Bintang jika rata-rata skor >= 90) ke SELURUH kru di batch tersebut.",
        "   • Gamification Engine memperbarui akumulasi bintang, level karir, membuka lencana prestasi, dan memperbarui peringkat leaderboard.",
        "5. Opsi B - REQUEST REVISION (Minta Perbaikan):",
        "   • Status misi berubah menjadi REVISION_REQUIRED.",
        "   • Head wajib mengisikan catatan instruksi revisi spesifik (misal: 'Mohon upload ulang bukti swab test bar').",
        "   • Notifikasi dikirimkan kembali ke Supervisor untuk dilakukan perbaikan audit."
    ]
    for step in p_head_steps:
        p = doc.add_paragraph(step)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== 4. MEKANISME GAMIFIKASI ====================
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. Mekanisme Gamifikasi & Star Economy")
    r_h4.font.color.rgb = COLOR_PRIMARY
    r_h4.font.bold = True

    p = doc.add_paragraph(
        "Sistem gamifikasi Re.juve dirancang untuk menumbuhkan motivasi intrinsik dan kebanggaan profesi bagi para Store Specialist:"
    )
    p.paragraph_format.space_after = Pt(8)

    gami_points = [
        ("Formula Bintang Misi: ", "Skor evaluasi $\\ge 90$ mendapatkan 5 Bintang Penuh (⭐⭐⭐⭐⭐); Skor 80–89 mendapatkan 4 Bintang; Skor 70–79 mendapatkan 3 Bintang."),
        ("Prinsip Store-Wide Multi-Crew Reward: ", "Karena operasional gerai adalah kerja tim, bintang dibagikan secara serentak ke seluruh kru aktif di gerai tersebut saat misi disetujui. Ini mendorong budaya saling mengingatkan SOP antar kru."),
        ("Tangga Karir 10 Tingkat (Level Progression): ", "Kru berkembang dari Level 1 (Trainee Specialist - 0 Bintang) hingga Level 10 (Master Store Leader - > 2.500 Bintang)."),
        ("Lencana Prestasi (Achievement Badges): ", "Diberikan saat mencapai milestone khusus, seperti 'Cold-Chain Master' (100% audit suhu sempurna), 'Speed Demon' (kecepatan barista terbaik), dan 'Hygiene Guardian'."),
        ("Leaderboard Real-Time: ", "Menampilkan kompetisi sehat peringkat individu teratas dan peringkat gerai terbaik secara transparan.")
    ]
    for g_title, g_desc in gami_points:
        gp = doc.add_paragraph(style='List Bullet')
        gp.paragraph_format.space_after = Pt(4)
        r_gt = gp.add_run(g_title)
        r_gt.bold = True
        r_gt.font.color.rgb = COLOR_PRIMARY
        gp.add_run(g_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== 5. ISOLASI DATA MULTI-TENANT ====================
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. Isolasi Data & Keamanan Multi-Tenant")
    r_h5.font.color.rgb = COLOR_PRIMARY
    r_h5.font.bold = True

    p = doc.add_paragraph(
        "Untuk menjamin kerahasiaan dan ketertiban audit antar wilayah, sistem menerapkan pemisahan data mutlak:"
    )
    p.paragraph_format.space_after = Pt(8)

    iso_table = doc.add_table(rows=6, cols=4)
    iso_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    iso_hdrs = ["Akun Pengujian", "Nama Pejabat", "Wilayah / Batch Penugasan", "Akses Data yang Terlihat"]
    for i, title in enumerate(iso_hdrs):
        iso_table.rows[0].cells[i].text = title
        set_cell_background(iso_table.rows[0].cells[i], COLOR_BG_HEADER)
        set_cell_margins(iso_table.rows[0].cells[i], top=100, bottom=100, left=120, right=120)
        p = iso_table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)

    iso_rows = [
        ("Supervisor 1", "Budi Santoso", "Batch 1 (Grand Indonesia) & Batch 2 (Senayan City)", "Hanya dapat menilai misi di Batch 1 & 2. Tidak dapat melihat Batch 3."),
        ("Supervisor 2", "Dewi Lestari", "Batch 3 (Pondok Indah Mall)", "Hanya dapat menilai misi di Batch 3. Terpisah dari Batch 1 & 2."),
        ("Head 1", "Ahmad Dahlan", "Batch 1 & Batch 2", "Antrean approval hanya memuat evaluasi dari Batch 1 & 2."),
        ("Head 2", "Citra Dewi", "Batch 3 (PIM)", "Antrean approval hanya memuat evaluasi dari Batch 3."),
        ("Superadmin", "Siti Rahmawati", "Seluruh Cabang Nasional", "Akses master ke semua batch, template, user, dan approval.")
    ]

    for row_idx, data in enumerate(iso_rows, start=1):
        row_cells = iso_table.rows[row_idx].cells
        bg_color = COLOR_BG_GRAY if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(9)
                if col_idx == 0:
                    r.font.bold = True
                    r.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== 6. PANDUAN SLIDE PRESENTASI ====================
    h6 = doc.add_heading(level=1)
    r_h6 = h6.add_run("6. Panduan 10-Slide Pitch Deck Presentasi Manajemen")
    r_h6.font.color.rgb = COLOR_PRIMARY
    r_h6.font.bold = True

    p = doc.add_paragraph(
        "Gunakan susunan 10 slide ini saat mempresentasikan sistem di depan Direksi / Tim Manajemen:"
    )
    p.paragraph_format.space_after = Pt(8)

    slides = [
        ("Slide 1: Judul & Pembuka", 
         "Re.juve Operational Excellence & Gamification Platform: Transformasi Kepatuhan SOP Menjadi Budaya Berprestasi."),
        ("Slide 2: Masalah Operasional Saat Ini", 
         "Audit kertas lambat, feedback revisi tidak terdokumentasi rapi, dan motivasi kru gerai sulit diukur secara objektif."),
        ("Slide 3: Solusi Siklus 3 Minggu (3-Week Batch Cycle)", 
         "Membagi audit operasional menjadi 3 tema mingguan: Week 1 (Sanitasi & Suhu), Week 2 (Kualitas & Layanan), Week 3 (Closing & Stok)."),
        ("Slide 4: Struktur Wewenang 4 Pilar (RBAC)", 
         "Crew (Pelaksana) -> Supervisor (Penilai Lapangan) -> Head (Pengambil Keputusan) -> Superadmin (Pengendali Master)."),
        ("Slide 5: Alur Evaluasi Lapangan Supervisor", 
         "Supervisor input skor per kriteria SOP + wajib melampirkan foto bukti observasi langsung dari gerai."),
        ("Slide 6: Quality Gate: Approval & Request Revision", 
         "Head meninjau bukti audit; APPROVE memicu bintang serentak, REVISE mengembalikan evaluasi dengan catatan wajib perbaikan."),
        ("Slide 7: Mesin Gamifikasi & Star Economy", 
         "Distribusi bintang store-wide, level karir 1–10, pembukaan lencana keahlian, dan ranking leaderboard nasional."),
        ("Slide 8: Master Template SOP Modular", 
         "Kemudahan Superadmin mengelola paket standar gerai (Mall, Kiosk, Onboarding) dan menerapkannya dalam 1 klik."),
        ("Slide 9: Keamanan Data & Isolasi Wilayah", 
         "Pemisahan data mutlak antar Supervisor dan Head per batch gerai untuk menjamin kerahasiaan dan fokus cabang."),
        ("Slide 10: Dampak Bisnis & Kesimpulan", 
         "Kepatuhan standar produk #CleanLabel 100% konsisten, efisiensi audit naik 60%, dan retensi/motivasi kru meningkat.")
    ]

    for s_title, s_desc in slides:
        sp = doc.add_paragraph(style='List Bullet')
        sp.paragraph_format.space_after = Pt(4)
        r_st = sp.add_run(f"{s_title}: ")
        r_st.bold = True
        r_st.font.color.rgb = COLOR_SECONDARY
        sp.add_run(s_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Footer note
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_foot.add_run("Tautan Live Demo: https://gamification-dde4b.web.app • Re.juve True Cold-Pressed Operations")
    r_f.font.size = Pt(9.5)
    r_f.font.italic = True
    r_f.font.color.rgb = COLOR_TEXT_MUTED

    # Save to file
    output_path = "/Users/ikhsan/Documents/dev/Gamification/Dokumen_Sistem_dan_Alur_Gamifikasi_Rejuve.docx"
    doc.save(output_path)
    print(f"File successfully created at: {output_path}")

if __name__ == "__main__":
    create_rejuve_doc()
