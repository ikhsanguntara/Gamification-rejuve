import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_system_doc():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette
    COLOR_PRIMARY = RGBColor(131, 24, 67)     # Re.juve Burgundy (#831843)
    COLOR_SECONDARY = RGBColor(157, 23, 77)   # Deep Rose (#9d174d)
    COLOR_TEXT_MAIN = RGBColor(30, 41, 59)    # Slate 800 (#1e293b)
    COLOR_TEXT_MUTED = RGBColor(100, 116, 139)# Slate 500 (#64748b)
    COLOR_CODE = RGBColor(15, 23, 42)         # Slate 900
    COLOR_BG_HEADER = "831843"                # Hex Burgundy
    COLOR_BG_LIGHT = "FDF2F8"                 # Light Pink
    COLOR_BG_GRAY = "F8FAFC"                  # Light Slate
    COLOR_BG_CODE = "F1F5F9"                  # Code block background

    # Set normal style font
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = COLOR_TEXT_MAIN

    # ==================== COVER / HEADER ====================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    r_brand = p_title.add_run("RE.JUVE DIGITAL ECOSYSTEM\n")
    r_brand.font.size = Pt(13)
    r_brand.font.bold = True
    r_brand.font.color.rgb = COLOR_SECONDARY

    r_title = p_title.add_run("Dokumen Spesifikasi Teknis & Arsitektur Sistem\n")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    r_sub = p_title.add_run("System Architecture, Data Schema, Multi-Store Orchestration & State Machine")
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = COLOR_TEXT_MUTED

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(20)
    r_meta = p_meta.add_run("Dokumen Spesifikasi Teknis Perangkat Lunak (Software Technical Specification) • Versi 2.0")
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = COLOR_TEXT_MUTED

    # Divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(16)
    r_div = p_div.add_run("―" * 55)
    r_div.font.color.rgb = RGBColor(226, 232, 240)
    r_div.font.size = Pt(14)

    # ==================== 1. SYSTEM OVERVIEW & TECH STACK ====================
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Gambaran Umum & Tumpukan Teknologi (Technology Stack)")
    r_h1.font.color.rgb = COLOR_PRIMARY
    r_h1.font.bold = True

    p = doc.add_paragraph(
        "Platform Gamifikasi & Manajemen Kepatuhan SOP Re.juve dibangun dengan arsitektur Single Page Application (SPA) "
        "modern yang mengutamakan kecepatan respons, reaktivitas instan antar peran pengguna, dan performa tinggi pada perangkat mobile/tablet."
    )
    p.paragraph_format.space_after = Pt(8)

    t_tech = doc.add_table(rows=7, cols=3)
    t_tech.alignment = WD_TABLE_ALIGNMENT.CENTER

    th = ["Komponen Lapisan", "Teknologi / Library", "Fungsi & Peran Utama"]
    for i, title in enumerate(th):
        t_tech.rows[0].cells[i].text = title
        set_cell_background(t_tech.rows[0].cells[i], COLOR_BG_HEADER)
        set_cell_margins(t_tech.rows[0].cells[i], top=90, bottom=90, left=100, right=100)
        p = t_tech.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)

    tech_rows = [
        ("Framework Utama", "Nuxt 3.15 + Vue 3.5", "Rendering modern, Composition API, SSR/SSG/SPA Hybrid capabilities, modular routing."),
        ("State Management", "Pinia 3.0 + Pinia/Nuxt", "Reactivity engine pusat, isolasi data, dan orkestrasi transaksi multi-store."),
        ("Styling & Design System", "Tailwind CSS 3.4", "Sistem utilitas CSS responsif, Re.juve Brand Palette (Burgundy & Rose), Dark/Light mode."),
        ("Ikonografi & Animasi", "Lucide Vue Next + Canvas-Confetti", "Visual interaktif, efek perayaan Star Burst, dan penanda status antarmuka."),
        ("Penyimpanan Sesi", "VueUse + LocalStorage Adapter", "Persistensi sesi login, identitas persona aktif, dan sinkronisasi lintas halaman."),
        ("Hosting & Infrastruktur", "Firebase Hosting + Nitro Engine", "Distribusi aset statis global, latensi rendah (< 100ms), enkripsi SSL otomatis.")
    ]

    for row_idx, data in enumerate(tech_rows, start=1):
        row_cells = t_tech.rows[row_idx].cells
        bg_color = COLOR_BG_GRAY if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=90, right=90)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(9)
                if col_idx == 0:
                    r.font.bold = True
                    r.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== 2. SKEMA DATA & MODEL ENTITAS (DATA SCHEMA) ====================
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. Skema Data & Model Entitas (Data Model Specification)")
    r_h2.font.color.rgb = COLOR_PRIMARY
    r_h2.font.bold = True

    p = doc.add_paragraph(
        "Sistem menggunakan struktur data terelasi yang terkelola secara reaktif dalam memori Pinia store:"
    )
    p.paragraph_format.space_after = Pt(8)

    entities = [
        ("1. Entitas BATCH (`batches`)", 
         "Representasi siklus 3 minggu satu cabang gerai. Memuat informasi lokasi gerai, penugasan Supervisor & Head, konfigurasi approval threshold, dan status aktif minggu berjalan."),
        ("2. Entitas USER (`users / userDirectory`)", 
         "Direktori pengguna lengkap dengan peran (CREW, SUPERVISOR, HEAD, SUPERADMIN), alamat email kerja, lokasi cabang bertugas, batch ID penugasan, serta saldo bintang dan level karir."),
        ("3. Entitas MISSION (`missions`)", 
         "Daftar tugas SOP mingguan per batch. Memuat kode misi (M-01 s/d M-12), kategori (Suhu Dingin, Sanitasi, Kualitas Buah, Layanan), kriteria persyaratan SOP, skor evaluasi, bukti foto, dan status siklus."),
        ("4. Entitas APPROVAL (`approvals`)", 
         "Antrean persetujuan yang diterima Head of Operations. Berisi ringkasan skor rata-rata, daftar nilai individu per kru, foto bukti audit lapangan, catatan revisi, dan timestamp keputusan."),
        ("5. Entitas MASTER TEMPLATE PACKAGE (`templatePackages`)", 
         "Katalog master SOP siap pakai (Paket Standar Mall 12 Misi, Kiosk Express 6 Misi, Onboarding 6 Misi) yang dapat di-clone dan diterapkan ke batch baru secara instan."),
        ("6. Entitas GAMIFIKASI & BADGES (`badges & starTransactions`)", 
         "Catatan riwayat pencairan bintang, kenaikan level, dan progress pembukaan lencana prestasi bagi setiap kru.")
    ]

    for e_title, e_desc in entities:
        ep = doc.add_paragraph(style='List Bullet')
        ep.paragraph_format.space_after = Pt(4)
        r_et = ep.add_run(f"{e_title}: ")
        r_et.bold = True
        r_et.font.color.rgb = COLOR_SECONDARY
        ep.add_run(e_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table of Data Schema Fields
    p_st = doc.add_paragraph()
    r_st = p_st.add_run("Spesifikasi Field Entitas Utama:")
    r_st.bold = True

    t_schema = doc.add_table(rows=8, cols=4)
    t_schema.alignment = WD_TABLE_ALIGNMENT.CENTER

    sh_hdrs = ["Entitas", "Primary Field", "Tipe Data", "Deskripsi & Relasi"]
    for i, title in enumerate(sh_hdrs):
        t_schema.rows[0].cells[i].text = title
        set_cell_background(t_schema.rows[0].cells[i], COLOR_BG_HEADER)
        set_cell_margins(t_schema.rows[0].cells[i], top=90, bottom=90, left=100, right=100)
        p = t_schema.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)

    schema_rows = [
        ("Batch", "id, code, assignment", "String, Object", "Relasi: assignment.supervisorId -> User, assignment.headId -> User, assignment.crewIds -> Array[User]."),
        ("User", "id, role, batchId", "String, Enum", "Relasi: batchId -> Batch.id. Role: CREW | SUPERVISOR | HEAD | SUPERADMIN."),
        ("Mission", "id, batchId, week, status", "String, Int, Enum", "Relasi: batchId -> Batch.id. Status: LOCKED | IN_PROGRESS | DRAFT | PENDING_REVIEW | COMPLETED | REVISION_REQUIRED."),
        ("CrewEvaluation", "crewId, score, stars", "String, Int, Int", "Sub-dokumen di Mission & Approval. Menyimpan nilai 0-100 per kru per misi."),
        ("ApprovalItem", "id, missionId, batchId", "String, String", "Relasi: missionId -> Mission.id, batchId -> Batch.id, supervisorId -> User.id."),
        ("TemplatePackage", "id, code, templates", "String, Array[Obj]", "Katalog master butir SOP per minggu yang dapat di-generate ke misi batch."),
        ("Badge", "id, category, progress", "String, Enum, Int", "Milestone gamifikasi kru (Cold Chain Master, Speed Demon, Hygiene Champion).")
    ]

    for row_idx, data in enumerate(schema_rows, start=1):
        row_cells = t_schema.rows[row_idx].cells
        bg_color = COLOR_BG_GRAY if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=90, right=90)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(9)
                if col_idx == 0:
                    r.font.bold = True
                    r.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== 3. ORKESTRASI MULTI-STORE ====================
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Arsitektur State Management & Orkestrasi Multi-Store")
    r_h3.font.color.rgb = COLOR_PRIMARY
    r_h3.font.bold = True

    p = doc.add_paragraph(
        "Sistem tidak bergantung pada monolithic store, melainkan membagi logika ke dalam 6 Pinia Store modular yang "
        "berkomunikasi secara atomik dan reaktif:"
    )
    p.paragraph_format.space_after = Pt(8)

    stores_info = [
        ("1. `userStore` (Autentikasi & Direktori Pengguna):", 
         "Menyimpan identitas sesi, daftar seluruh pejabat/kru, dan menyediakan getter wewenang (isCrew, isSupervisor, isHead, isSuperadmin)."),
        ("2. `batchStore` (Manajemen Batch & Gerai):", 
         "Menyediakan getter `accessibleBatches` yang menyaring daftar gerai sesuai penugasan pengguna yang sedang aktif."),
        ("3. `missionStore` (Manajemen Siklus & Misi Gerai):", 
         "Menyimpan butir SOP, status per minggu, nilai per kru, dan menangani transisi status misi."),
        ("4. `evaluationStore` (Workspace Penilaian Supervisor):", 
         "Menyimpan lembar kerja sementara, perhitungan skor rata-rata otomatis, unggahan bukti foto, dan validasi kelengkapan sebelum submit."),
        ("5. `approvalStore` (Quality Gate & Antrean Persetujuan Head):", 
         "Menyediakan getter `userApprovals` yang hanya memuat antrean sesuai wilayah Head, aksi approve (memicu reward), dan aksi revise (mewajibkan feedback)."),
        ("6. `gamificationStore` (Mesin Reward, Star Economy & Leaderboard):", 
         "Mengatur pencairan bintang store-wide, akumulasi XP, kalkulasi level up, pembukaan badge, dan peringkat leaderboard antar-gerai.")
    ]

    for s_title, s_desc in stores_info:
        sp = doc.add_paragraph(style='List Bullet')
        sp.paragraph_format.space_after = Pt(4)
        r_st = sp.add_run(f"{s_title} ")
        r_st.bold = True
        r_st.font.color.rgb = COLOR_PRIMARY
        sp.add_run(s_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== 4. STATE MACHINE MISI & APPROVAL ====================
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. State Machine & Transisi Status Siklus Misi")
    r_h4.font.color.rgb = COLOR_PRIMARY
    r_h4.font.bold = True

    p = doc.add_paragraph(
        "Setiap butir misi SOP di dalam siklus batch mengalami transisi status yang terstandarisasi:"
    )
    p.paragraph_format.space_after = Pt(8)

    states = [
        ("1. LOCKED: ", "Misi di minggu mendatang yang belum dibuka (misal: Week 3 saat batch masih di Week 2). Tidak dapat dinilai."),
        ("2. IN_PROGRESS / ACTIVE: ", "Misi di minggu aktif yang siap dijalankan kru dan diobservasi oleh Supervisor."),
        ("3. DRAFT: ", "Supervisor telah menginput nilai atau mengunggah bukti, namun masih disimpan sementara di perangkat lokal."),
        ("4. PENDING_REVIEW: ", "Supervisor telah mengirimkan evaluasi lengkap. Misi berpindah ke antrean Head of Operations untuk ditinjau."),
        ("5. APPROVED / COMPLETED: ", "Head of Operations menyetujui evaluasi. Sistem secara atomik mencairkan bintang (+5 Stars) ke seluruh kru dan menutup misi."),
        ("6. REVISION_REQUIRED: ", "Head of Operations menolak evaluasi dan mewajibkan perbaikan. Misi dikembalikan ke Supervisor dengan catatan instruksi.")
    ]

    for st_title, st_desc in states:
        stp = doc.add_paragraph(style='List Bullet')
        stp.paragraph_format.space_after = Pt(4)
        r_s = stp.add_run(st_title)
        r_s.bold = True
        r_s.font.color.rgb = COLOR_SECONDARY
        stp.add_run(st_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== 5. TRANSAKSI ATOMIK SAAT APPROVAL ====================
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. Logika Transaksi Atomik Saat Persetujuan Head (Atomic Reward Minting)")
    r_h5.font.color.rgb = COLOR_PRIMARY
    r_h5.font.bold = True

    p = doc.add_paragraph(
        "Ketika Head of Operations menekan tombol 'Approve Evaluation', sistem menjalankan transaksi atomik 5 langkah simultan "
        "untuk menjamin konsistensi data gamifikasi:"
    )
    p.paragraph_format.space_after = Pt(8)

    steps = [
        ("Langkah 1: Update Status Misi & Approval: ", "Status Approval berubah menjadi 'APPROVED' dan status Mission di missionStore berubah menjadi 'COMPLETED'."),
        ("Langkah 2: Perhitungan Bintang Terkalkulasi: ", "Jika skor rata-rata $\\ge 90$, sistem mengalokasikan 5 bintang penuh per kru aktif. Jika 80–89, alokasi 4 bintang."),
        ("Langkah 3: Minting Bintang ke Seluruh Kru Gerai: ", "Saldo bintang setiap kru di userStore dan gamificationStore bertambah serentak (Store-Wide Reward)."),
        ("Langkah 4: Evaluasi Level Up & Badge Trigger: ", "Sistem mengevaluasi apakah akumulasi bintang kru melewati ambang level berikutnya (Level 1–10) atau memenuhi syarat lencana (Badge Unlock)."),
        ("Langkah 5: Pembaruan Leaderboard & Metrik Batch: ", "Peringkat kru di leaderboard dan metrik kelulusan batch (completion rate & total stars) diperbarui secara real-time.")
    ]

    for s_title, s_desc in steps:
        stp = doc.add_paragraph(style='List Bullet')
        stp.paragraph_format.space_after = Pt(4)
        r_st = stp.add_run(s_title)
        r_st.bold = True
        r_st.font.color.rgb = COLOR_PRIMARY
        stp.add_run(s_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== 6. ISOLASI DATA MULTI-TENANT & KEAMANAN ====================
    h6 = doc.add_heading(level=1)
    r_h6 = h6.add_run("6. Mekanisme Isolasi Data Multi-Tenant & Keamanan Akses")
    r_h6.font.color.rgb = COLOR_PRIMARY
    r_h6.font.bold = True

    p = doc.add_paragraph(
        "Sistem menerapkan aturan isolasi data ketat di sisi logika aplikasi untuk mencegah kebocoran data antar cabang:"
    )
    p.paragraph_format.space_after = Pt(8)

    sec_rules = [
        ("Aturan 1 (Supervisor Isolation): ", "Getter `accessibleBatches` memfilter batch di mana `assignment.supervisorId === userStore.currentUserId`. Supervisor hanya dapat melihat dan menilai batch yang ditugaskan kepadanya."),
        ("Aturan 2 (Head Approval Isolation): ", "Getter `userApprovals` hanya memuat antrean approval dari batch di mana `assignment.headId === userStore.currentUserId`."),
        ("Aturan 3 (Crew Scope Isolation): ", "Kru gerai secara otomatis terkunci pada `currentUser.batchId` dan tidak memiliki hak akses ke modul penilaian atau konfigurasi master."),
        ("Aturan 4 (Superadmin Master Access): ", "Superadmin memiliki akses bypass ke seluruh batch, template, direktori user, dan konfigurasi persetujuan.")
    ]

    for sr_title, sr_desc in sec_rules:
        srp = doc.add_paragraph(style='List Bullet')
        srp.paragraph_format.space_after = Pt(4)
        r_sr = srp.add_run(sr_title)
        r_sr.bold = True
        r_sr.font.color.rgb = COLOR_SECONDARY
        srp.add_run(sr_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Footer
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_foot.add_run("Re.juve Store Operations Gamification • Software Technical Specification Document")
    r_f.font.size = Pt(9)
    r_f.font.italic = True
    r_f.font.color.rgb = COLOR_TEXT_MUTED

    # Save to file
    output_path = "/Users/ikhsan/Documents/dev/Gamification/Dokumen_Spesifikasi_Teknis_Sistem_Rejuve.docx"
    doc.save(output_path)
    print(f"System Technical Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_system_doc()
