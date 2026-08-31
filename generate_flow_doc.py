import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_flow_doc():
    doc = Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    COLOR_PRIMARY = RGBColor(131, 24, 67)     # Burgundy (#831843)
    COLOR_SECONDARY = RGBColor(157, 23, 77)   # Deep Rose (#9d174d)
    COLOR_TEXT_MAIN = RGBColor(30, 41, 59)    # Slate 800 (#1e293b)
    COLOR_TEXT_MUTED = RGBColor(100, 116, 139)# Slate 500 (#64748b)
    COLOR_EMERALD = RGBColor(16, 185, 129)    # Emerald Green
    COLOR_AMBER = RGBColor(217, 119, 6)       # Amber
    COLOR_ROSE = RGBColor(225, 29, 72)        # Rose Red

    COLOR_BG_HEADER = "831843"
    COLOR_BG_GRAY = "F8FAFC"
    COLOR_BG_STEP = "FDF2F8"

    # Set normal style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = COLOR_TEXT_MAIN

    # ==================== COVER / HEADER ====================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    r_brand = p_title.add_run("RE.JUVE OPERATIONAL EXCELLENCE\n")
    r_brand.font.size = Pt(13)
    r_brand.font.bold = True
    r_brand.font.color.rgb = COLOR_SECONDARY

    r_title = p_title.add_run("Buku Panduan Lengkap Alur & Flow Sistem\n")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    r_sub = p_title.add_run("End-to-End Process Flows, Decision Trees, Role Workflows, & Gamification Engine")
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = COLOR_TEXT_MUTED

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(20)
    r_meta = p_meta.add_run("Dokumen Alur Sistem Resmi • Versi 2.0 (Production Ready)")
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = COLOR_TEXT_MUTED

    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(16)
    r_div = p_div.add_run("―" * 55)
    r_div.font.color.rgb = RGBColor(226, 232, 240)
    r_div.font.size = Pt(14)

    # ==================== FLOW 1: END-TO-END SWIMLANE FLOW ====================
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Alur Utama End-to-End Sistem (High-Level Process Flow)")
    r_h1.font.color.rgb = COLOR_PRIMARY
    r_h1.font.bold = True

    p = doc.add_paragraph(
        "Alur operasional sistem berjalan secara berkesinambungan melibatkan 4 aktor utama. "
        "Berikut adalah urutan proses dari inisialisasi awal hingga pencairan reward bintang:"
    )
    p.paragraph_format.space_after = Pt(8)

    flow1_table = doc.add_table(rows=6, cols=3)
    flow1_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    th = ["Fase / Tahapan", "Aktor Utama", "Aktivitas & Output Sistem"]
    for i, title in enumerate(th):
        flow1_table.rows[0].cells[i].text = title
        set_cell_background(flow1_table.rows[0].cells[i], COLOR_BG_HEADER)
        set_cell_margins(flow1_table.rows[0].cells[i], top=90, bottom=90, left=100, right=100)
        p = flow1_table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)

    f1_rows = [
        ("Fase 1: Inisialisasi Batch", "👑 Superadmin", "Membuat batch gerai baru, memilih paket template SOP, menugaskan Supervisor & Head, serta memilih kru. Sistem meng-generate 12 misi untuk siklus 3 minggu."),
        ("Fase 2: Eksekusi SOP", "👤 Store Crew", "Menjalankan standar operasional di gerai (catat suhu chiller 2-4°C, sanitasi bar, uji rasa buah, speed of service barista)."),
        ("Fase 3: Audit Lapangan", "📋 Area Supervisor", "Mengunjungi gerai, menilai setiap butir kriteria SOP, menginput skor per kru, melampirkan foto bukti observasi, lalu mengirim evaluasi (Status: PENDING_REVIEW)."),
        ("Fase 4: Review Mutu", "🛡️ Head of Ops", "Memeriksa kesesuaian foto bukti dan skor. Mengambil keputusan: APPROVE (mencairkan bintang) atau REQUEST REVISION (mewajibkan perbaikan)."),
        ("Fase 5: Reward & Ranking", "⭐ Gamification Engine", "Pencairan otomatis +5 Bintang ke seluruh kru batch, kalkulasi kenaikan level (Level 1-10), pembukaan badge, dan pembaruan peringkat leaderboard.")
    ]

    for row_idx, data in enumerate(f1_rows, start=1):
        row_cells = flow1_table.rows[row_idx].cells
        bg_color = COLOR_BG_GRAY if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=90, right=90)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(9)
                if col_idx == 0:
                    r.font.bold = True
                    r.font.color.rgb = COLOR_PRIMARY
                elif col_idx == 1:
                    r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== FLOW 2: SUPERVISOR FIELD EVALUATION ====================
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. Alur Detil Penilaian Lapangan oleh Supervisor")
    r_h2.font.color.rgb = COLOR_PRIMARY
    r_h2.font.bold = True

    p = doc.add_paragraph(
        "Alur ini menjelaskan langkah kerja Supervisor saat melakukan audit langsung di gerai Re.juve:"
    )
    p.paragraph_format.space_after = Pt(6)

    spv_steps = [
        ("Langkah 2.1: Login & Seleksi Wilayah Otomatis",
         "Supervisor login ke sistem. Sistem secara otomatis menyaring dan hanya menampilkan batch gerai yang berada di bawah tanggung jawab Supervisor tersebut (misal: Budi Santoso hanya melihat Batch 1 GI & Batch 2 Senayan)."),
        ("Langkah 2.2: Buka Workspace Evaluasi (/evaluations)",
         "Supervisor memilih tab minggu aktif (Week 1, Week 2, atau Week 3) dan memilih butir misi SOP yang akan diaudit (contoh: 'Cek Suhu Chiller 2-4°C')."),
        ("Langkah 2.3: Pemeriksaan Butir Kriteria & Skoring Kru",
         "Supervisor mencocokkan kondisi fisik gerai dengan daftar checklist SOP. Supervisor memasukkan skor evaluasi (skala 0–100) untuk masing-masing kru aktif yang terdaftar di gerai tersebut."),
        ("Langkah 2.4: Wajib Melampirkan Bukti Foto (Photo Evidence)",
         "Supervisor wajib mengunggah minimal 1 foto bukti observasi langsung (misal: foto termometer display, lembar logbook suhu, atau foto bar press juice yang higienis)."),
        ("Langkah 2.5: Input Catatan & Komentar Evaluator",
         "Supervisor menuliskan ringkasan hasil observasi (misal: 'Suhu chiller stabil di 3.2°C, seluruh kru memakai sarung tangan lengkap')."),
        ("Langkah 2.6: Percabangan Tindakan (Simpan Draft vs Submit Head)",
         "• Opsi SIMPAN DRAFT: Evaluasi disimpan secara lokal dan dapat diedit kembali sewaktu-waktu.\n• Opsi KIRIM KE HEAD: Sistem memvalidasi kelengkapan foto & nilai, lalu mengubah status misi menjadi PENDING_REVIEW dan meneruskannya ke antrean Head of Operations.")
    ]

    for s_title, s_desc in spv_steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(s_title + "\n")
        r_t.bold = True
        r_t.font.color.rgb = COLOR_SECONDARY
        r_d = p.add_run(s_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== FLOW 3: HEAD APPROVAL & REVISION ====================
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Alur Detil Persetujuan & Keputusan oleh Head of Operations")
    r_h3.font.color.rgb = COLOR_PRIMARY
    r_h3.font.bold = True

    p = doc.add_paragraph(
        "Head of Operations bertindak sebagai gerbang kendali mutu akhir (*Quality Gate*). "
        "Berikut adalah alur peninjauan dan logika percabangan keputusan:"
    )
    p.paragraph_format.space_after = Pt(6)

    head_steps = [
        ("Langkah 3.1: Akses Antrean Persetujuan Terfilter (/approvals)",
         "Head of Operations membuka halaman antrean. Sistem hanya menampilkan pengajuan evaluasi dari gerai yang ditugaskan kepada Head tersebut."),
        ("Langkah 3.2: Inspeksi Bukti Audit Lapangan",
         "Head membuka detail evaluasi, memeriksa skor rata-rata gerai, nilai per kru, catatan Supervisor, dan memperbesar foto bukti audit untuk verifikasi standar Re.juve."),
        ("Langkah 3.3: Titik Keputusan A - Persetujuan (APPROVE)",
         "Jika hasil audit memenuhi standar Re.juve:\n"
         "1. Head menekan tombol 'Approve Store Evaluation'.\n"
         "2. Status Misi & Approval berubah menjadi COMPLETED / APPROVED.\n"
         "3. Gamification Engine secara otomatis mencairkan +5 Bintang (jika nilai >= 90) ke SELURUH kru di batch tersebut secara serentak.\n"
         "4. Saldo bintang kru bertambah, level karir naik jika memenuhi syarat, badge terbuka, dan ranking leaderboard diperbarui."),
        ("Langkah 3.4: Titik Keputusan B - Permintaan Perbaikan (REQUEST REVISION)",
         "Jika hasil audit belum memuaskan atau bukti foto kurang jelas:\n"
         "1. Head menekan tombol 'Request Revision'.\n"
         "2. Head WAJIB mengisikan catatan perbaikan spesifik (misal: 'Foto logbook suhu buram, mohon upload ulang lembar logbook sore').\n"
         "3. Status Misi berubah menjadi REVISION_REQUIRED.\n"
         "4. Notifikasi revisi otomatis terkirim kembali ke Supervisor untuk dilakukan perbaikan audit.")
    ]

    for s_title, s_desc in head_steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(s_title + "\n")
        r_t.bold = True
        r_t.font.color.rgb = COLOR_PRIMARY
        r_d = p.add_run(s_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== FLOW 4: RESUBMISSION ====================
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. Alur Perbaikan Revisi oleh Supervisor (Revision Cycle)")
    r_h4.font.color.rgb = COLOR_PRIMARY
    r_h4.font.bold = True

    p = doc.add_paragraph(
        "Ketika evaluasi ditolak dengan status REVISION_REQUIRED, siklus perbaikan berjalan sebagai berikut:"
    )
    p.paragraph_format.space_after = Pt(6)

    rev_steps = [
        ("1. Notifikasi Revisi Diterima: ", "Supervisor melihat badge '1 Rev' di menu navigasi Evaluasi dan notifikasi permintaan perbaikan dari Head."),
        ("2. Membaca Catatan Head: ", "Supervisor membuka misi terkait dan membaca kotak instruksi revisi berwarna kuning (contoh: instruksi melengkapi foto ATP swab test)."),
        ("3. Melakukan Perbaikan Lapangan: ", "Supervisor memperbarui nilai atau mengunggah foto bukti tambahan yang diminta."),
        ("4. Mengirim Ulang (Resubmit): ", "Supervisor menekan 'Kirim Ulang Evaluasi'. Status misi kembali menjadi PENDING_REVIEW dan masuk kembali ke antrean Head.")
    ]
    for r_title, r_desc in rev_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_rt = p.add_run(r_title)
        r_rt.bold = True
        r_rt.font.color.rgb = COLOR_AMBER
        p.add_run(r_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== FLOW 5: CREW GAMIFICATION JOURNEY ====================
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. Alur Pengalaman Kru Gerai (Store Crew Gamification Journey)")
    r_h5.font.color.rgb = COLOR_PRIMARY
    r_h5.font.bold = True

    p = doc.add_paragraph(
        "Dari sudut pandang Store Crew, sistem memberikan umpan balik positif secara berkala:"
    )
    p.paragraph_format.space_after = Pt(6)

    crew_steps = [
        ("1. Pantau Misi Gerai (/missions & /journey): ", "Kru melihat misi aktif minggu ini beserta standar checklist yang harus dijaga di gerai."),
        ("2. Penerimaan Reward Instan: ", "Saat Head menyetujui misi gerai, kru menerima notifikasi perolehan +5 Bintang."),
        ("3. Efek Animasi & Perayaan: ", "Tampil animasi Star Burst dan notifikasi selebrasi di dashboard."),
        ("4. Kenaikan Level Karir: ", "Akumulasi bintang mendorong level kru dari Level 1 (Trainee) hingga Level 10 (Master Leader)."),
        ("5. Pembukaan Lencana (Badges): ", "Membuka lencana keahlian seperti 'Cold-Chain Master' dan 'Hygiene Guardian'."),
        ("6. Kompetisi Leaderboard: ", "Melihat peringkat kontribusi kru dan peringkat gerai di Leaderboard nasional.")
    ]
    for c_title, c_desc in crew_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_ct = p.add_run(c_title)
        r_ct.bold = True
        r_ct.font.color.rgb = COLOR_EMERALD
        p.add_run(c_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ==================== FLOW 6: ADMIN MASTER TEMPLATE & BATCH ====================
    h6 = doc.add_heading(level=1)
    r_h6 = h6.add_run("6. Alur Pengelolaan Master Template & Batch oleh Superadmin")
    r_h6.font.color.rgb = COLOR_PRIMARY
    r_h6.font.bold = True

    p = doc.add_paragraph(
        "Superadmin mengelola standardisasi SOP nasional dan inisialisasi batch baru:"
    )
    p.paragraph_format.space_after = Pt(6)

    admin_steps = [
        ("Alur Master Template SOP (/admin/templates): ", 
         "Superadmin membuat Paket Master (misal: Standar Flagship Mall 12 Misi) -> Menambah butir SOP per minggu -> Melakukan duplikasi 1-klik jika ada penyesuaian untuk format gerai Kiosk -> Menerapkan paket ke batch tujuan."),
        ("Alur Pembuatan Batch Gerai (/admin/batches/create): ", 
         "Superadmin memasukkan Nama & Lokasi Gerai -> Memilih Supervisor & Head penanggung jawab -> Memilih Paket SOP -> Memilih roster kru via tag checklist -> Menyimpan batch. Sistem langsung mengaktifkan gerai di siklus Week 1.")
    ]
    for a_title, a_desc in admin_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_at = p.add_run(a_title)
        r_at.bold = True
        r_at.font.color.rgb = COLOR_PRIMARY
        p.add_run(a_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Footer
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_foot.add_run("Re.juve True Cold-Pressed Operations • End-to-End System Flow Guide")
    r_f.font.size = Pt(9)
    r_f.font.italic = True
    r_f.font.color.rgb = COLOR_TEXT_MUTED

    output_path = "/Users/ikhsan/Documents/dev/Gamification/Dokumen_Alur_dan_Flow_Sistem_Rejuve.docx"
    doc.save(output_path)
    print(f"System Flow Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_flow_doc()
